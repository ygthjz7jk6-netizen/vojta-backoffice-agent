from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path.cwd()
DEMO = ROOT / "demo-data"
FILES = DEMO / "files"
CALENDAR = DEMO / "google-calendar"


def ensure_dirs() -> None:
    for path in [
        FILES / "excel",
        FILES / "pdf",
        FILES / "docx",
        FILES / "contracts",
        CALENDAR,
    ]:
        path.mkdir(parents=True, exist_ok=True)


def style_sheet(ws, widths: dict[int, int]) -> None:
    header_fill = PatternFill("solid", fgColor="1F4E46")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def create_excel_files() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Pipeline"
    ws.append(["ID", "Nemovitost", "Lokalita", "Faze", "Cena CZK", "Pravdepodobnost", "Broker", "Dalsi krok"])
    rows = [
        ["D-9003", "Byt Vinohrady Korunni", "Praha 2 - Vinohrady", "rezervacni smlouva", 12350000, 0.85, "Marie Vankova", "kontrola advokatni uschovy"],
        ["D-9004", "Byt 3+1, V Haji", "Praha 7 - Holesovice", "rezervacni smlouva", 11650000, 0.78, "Pepa Novak", "doplneni rekonstrukce jadra"],
        ["D-9005", "Byt 2+kk, Ortenovo namesti", "Praha 7 - Holesovice", "v jednani", 8620000, 0.55, "Pepa Novak", "potvrdit financovani"],
        ["D-9006", "Loft Karlin", "Praha 8 - Karlin", "v jednani", 18100000, 0.45, "Marie Vankova", "dodat odhad najmu"],
        ["D-9007", "Cinzak Tusarova", "Praha 7 - Holesovice", "akvizice", 84500000, 0.32, "Robert Linhart", "vyzadat PENB"],
    ]
    for row in rows:
        ws.append(row)
    style_sheet(ws, {1: 12, 2: 28, 3: 22, 4: 22, 5: 14, 6: 18, 7: 18, 8: 32})
    for row in ws.iter_rows(min_row=2, min_col=5, max_col=5):
        row[0].number_format = '#,##0 "Kc"'
    for row in ws.iter_rows(min_row=2, min_col=6, max_col=6):
        row[0].number_format = "0%"
    wb.save(FILES / "excel" / "pipeline_obchodu_kveten_2026.xlsx")

    wb = Workbook()
    ws = wb.active
    ws.title = "Mesicni KPI"
    ws.append(["Mesic", "Leady", "Kvalifikovane leady", "Prohlidky", "Nove nemovitosti", "Prodano"])
    metrics = [
        ["2025-11", 38, 17, 12, 5, 1],
        ["2025-12", 31, 15, 10, 4, 1],
        ["2026-01", 44, 22, 17, 6, 0],
        ["2026-02", 52, 28, 21, 8, 1],
        ["2026-03", 61, 33, 25, 7, 1],
        ["2026-04", 58, 31, 24, 9, 2],
    ]
    for row in metrics:
        ws.append(row)
    style_sheet(ws, {1: 14, 2: 12, 3: 20, 4: 14, 5: 18, 6: 12})
    line = LineChart()
    line.title = "Vyvoj leadu"
    line.y_axis.title = "Pocet"
    line.x_axis.title = "Mesic"
    line.add_data(Reference(ws, min_col=2, min_row=1, max_row=7), titles_from_data=True)
    line.set_categories(Reference(ws, min_col=1, min_row=2, max_row=7))
    ws.add_chart(line, "H2")
    bar = BarChart()
    bar.title = "Prodane nemovitosti"
    bar.add_data(Reference(ws, min_col=6, min_row=1, max_row=7), titles_from_data=True)
    bar.set_categories(Reference(ws, min_col=1, min_row=2, max_row=7))
    ws.add_chart(bar, "H18")
    wb.save(FILES / "excel" / "mesicni_report_leady_prodeje.xlsx")

    wb = Workbook()
    ws = wb.active
    ws.title = "Data quality"
    ws.append(["Property ID", "Nemovitost", "Problem", "Priorita", "Doporuceny dalsi krok"])
    rows = [
        ["P-1002", "Byt 3+1, V Haji", "chybi rozsah rekonstrukce jadra", "vysoka", "vyzadat technicky popis od majitele"],
        ["P-1004", "Cinzak Tusarova", "chybi stavebni upravy a PENB", "vysoka", "zadat seznam uprav a objednat PENB"],
        ["P-1006", "Loft Karlin", "chybi informace o zmene dispozice", "stredni", "overit dokumentaci SVJ"],
        ["P-1009", "Nebytovy prostor Smichov", "chybi rekonstrukce vzduchotechniky", "stredni", "vyzadat revizni zpravu"],
        ["P-1011", "Pozemek Kladno", "chybi UPI", "vysoka", "podat zadost o uzemne planovaci informaci"],
    ]
    for row in rows:
        ws.append(row)
    style_sheet(ws, {1: 14, 2: 28, 3: 34, 4: 12, 5: 42})
    wb.save(FILES / "excel" / "data_quality_chybejici_podklady.xlsx")


def create_docx_files() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    title = doc.add_heading("Týdenní report pro vedení", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Období: 27. 4. 2026 - 3. 5. 2026 | Připravil: Pepa Novák")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_heading("Executive summary", level=1)
    for text in [
        "Přišlo 14 nových leadů, z toho 8 bylo kvalifikováno jako obchodně relevantních.",
        "Proběhlo 6 prohlídek, dvě z nich jsou druhé prohlídky s vysokou pravděpodobností nabídky.",
        "Největší riziko týdne jsou chybějící technické podklady u nemovitostí P-1002 a P-1004.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Doporučené kroky", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Priorita"
    hdr[1].text = "Oblast"
    hdr[2].text = "Akce"
    hdr[3].text = "Owner"
    for row in [
        ["1", "Data quality", "Doplnit rekonstrukce a PENB u rizikových nemovitostí.", "Pepa"],
        ["2", "Obchod", "Potvrdit financování u P-1001 a druhou prohlídku P-1010.", "Broker tým"],
        ["3", "Monitoring", "Každé ráno posílat benchmark Holešovic.", "Agent"],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.save(FILES / "docx" / "tydenni_report_vedeni_2026-05-03.docx")

    doc = Document()
    doc.add_heading("Zápis z akviziční porady", 0)
    doc.add_paragraph("Datum: 30. 4. 2026 | Účastníci: Pepa Novák, Marie Vaňková, Robert Linhart")
    doc.add_heading("Nemovitosti k řešení", level=1)
    for text in [
        "P-1004 Cinžák Tusarova: čekáme na PENB, seznam stavebních úprav a potvrzení výnosů z nájmů.",
        "P-1011 Pozemek Kladno: nutné ověřit územní plán a limity zastavitelnosti.",
        "P-1006 Loft Karlín: doplnit dokumentaci ke změně dispozice.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Rozhodnutí", level=1)
    doc.add_paragraph("Agent má každý den připravit seznam blokovaných nemovitostí a doporučený další krok.")
    doc.save(FILES / "docx" / "zapis_akvizicni_porada_2026-04-30.docx")


def create_pdf(path: Path, title: str, rows: list[list[str]], intro: str) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    story = [Paragraph(title, styles["Title"]), Spacer(1, 8), Paragraph(intro, styles["BodyText"]), Spacer(1, 12)]
    table = Table(rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E46")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB7B2")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7F6")]),
            ]
        )
    )
    story.append(table)
    doc.build(story)


def create_pdf_files() -> None:
    create_pdf(
        FILES / "pdf" / "penb_p1004_chybejici_vyzva.pdf",
        "Výzva k doložení PENB",
        [
            ["Pole", "Hodnota"],
            ["Nemovitost", "P-1004 Cinžák Tusarova, Praha 7 - Holešovice"],
            ["Stav", "PENB nebyl dodán"],
            ["Dopad", "Nelze publikovat kompletní inzertní podklady bez rizika compliance chyby"],
            ["Další krok", "Vyžádat PENB od majitele nebo objednat zpracování"],
        ],
        "Mock PDF dokument simulující přílohu z e-mailu nebo složky na Google Drive.",
    )
    create_pdf(
        FILES / "pdf" / "list_vlastnictvi_p1001_mock.pdf",
        "List vlastnictví - zkrácený opis",
        [
            ["Pole", "Hodnota"],
            ["Nemovitost", "P-1001 Byt 2+kk, Ortenovo náměstí"],
            ["Vlastník", "Klient C-0002"],
            ["Omezení vlastnického práva", "Bez záznamu v mock datech"],
            ["Poznámka", "Ověřit soulad výměry s prohlášením vlastníka"],
        ],
        "Nejedná se o skutečný výpis z KN, pouze realistická testovací napodobenina.",
    )
    create_pdf(
        FILES / "contracts" / "rezervacni_smlouva_p1002_mock.pdf",
        "Rezervační smlouva - pracovní verze",
        [
            ["Článek", "Obsah"],
            ["Předmět", "Rezervace bytu P-1002 Byt 3+1, V Háji"],
            ["Cena", "11 650 000 Kč"],
            ["Rezervační poplatek", "350 000 Kč"],
            ["Otevřené body", "Doplnit rozsah rekonstrukce bytového jádra a datum posledních stavebních úprav"],
        ],
        "Mock smlouva pro testování extrakce závazků, otevřených bodů a rizik.",
    )


def create_google_calendar_files() -> None:
    events = [
        {
            "id": "gcal_evt_20260506_0900_standup",
            "summary": "Interní standup obchodu",
            "start": {"dateTime": "2026-05-06T09:00:00+02:00", "timeZone": "Europe/Prague"},
            "end": {"dateTime": "2026-05-06T09:45:00+02:00", "timeZone": "Europe/Prague"},
            "status": "confirmed",
            "transparency": "opaque",
            "attendees": [{"email": "pepa@example.cz", "self": True}],
        },
        {
            "id": "gcal_evt_20260506_1430_free",
            "summary": "Volný blok pro prohlídku",
            "start": {"dateTime": "2026-05-06T14:30:00+02:00", "timeZone": "Europe/Prague"},
            "end": {"dateTime": "2026-05-06T15:30:00+02:00", "timeZone": "Europe/Prague"},
            "status": "confirmed",
            "transparency": "transparent",
            "attendees": [{"email": "pepa@example.cz", "self": True}],
        },
        {
            "id": "gcal_evt_20260507_1100_free",
            "summary": "Volný blok pro prohlídku",
            "start": {"dateTime": "2026-05-07T11:00:00+02:00", "timeZone": "Europe/Prague"},
            "end": {"dateTime": "2026-05-07T12:00:00+02:00", "timeZone": "Europe/Prague"},
            "status": "confirmed",
            "transparency": "transparent",
            "attendees": [{"email": "pepa@example.cz", "self": True}],
        },
        {
            "id": "gcal_evt_20260508_1000_management",
            "summary": "Porada s vedením",
            "start": {"dateTime": "2026-05-08T10:00:00+02:00", "timeZone": "Europe/Prague"},
            "end": {"dateTime": "2026-05-08T11:30:00+02:00", "timeZone": "Europe/Prague"},
            "status": "confirmed",
            "transparency": "opaque",
            "attendees": [{"email": "pepa@example.cz", "self": True}, {"email": "vedeni@example.cz"}],
        },
    ]
    payload = {
        "kind": "calendar#events",
        "summary": "Pepa Novak - pracovní kalendář",
        "timeZone": "Europe/Prague",
        "updated": datetime.now(timezone.utc).isoformat(),
        "items": events,
    }
    (CALENDAR / "google_calendar_events_mock.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    ics_lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Demo Backoffice Agent//Google Calendar Mock//CZ",
    ]
    for event in events:
        start = event["start"]["dateTime"].replace("-", "").replace(":", "").replace("+0200", "")
        end = event["end"]["dateTime"].replace("-", "").replace(":", "").replace("+0200", "")
        start = start.split("+")[0]
        end = end.split("+")[0]
        ics_lines.extend(
            [
                "BEGIN:VEVENT",
                f"UID:{event['id']}@demo.local",
                f"DTSTART;TZID=Europe/Prague:{start}",
                f"DTEND;TZID=Europe/Prague:{end}",
                f"SUMMARY:{event['summary']}",
                f"TRANSP:{'TRANSPARENT' if event['transparency'] == 'transparent' else 'OPAQUE'}",
                "END:VEVENT",
            ]
        )
    ics_lines.append("END:VCALENDAR")
    (CALENDAR / "pepa_google_calendar_mock.ics").write_text("\n".join(ics_lines) + "\n", encoding="utf-8")


def update_readme() -> None:
    readme = DEMO / "README.md"
    original = readme.read_text(encoding="utf-8")
    addition = """

## Napodobeniny reálných souborů

Kromě CSV exportů je ve složce `files/` druhá vrstva dat, která víc odpovídá tomu, co Pepa reálně dostává:

- `files/excel/`: Excel pipeline, měsíční KPI report a data-quality seznam.
- `files/pdf/`: mock PDF výzvy, list vlastnictví a pracovní rezervační smlouva.
- `files/docx/`: týdenní report a zápis z akviziční porady.
- `files/contracts/`: smluvní PDF s otevřenými body.
- `google-calendar/`: JSON napodobující odpověď Google Calendar API a `.ics` export.

Prakticky by agent nejdřív četl soubory z Google Drive/Gmailu, extrahoval strukturované informace a ukládal je do indexu. CSV v této složce proto reprezentují už očištěnou analytickou vrstvu, zatímco PDF/DOCX/XLSX ukazují původní pracovní realitu.
"""
    if "## Napodobeniny reálných souborů" not in original:
        readme.write_text(original.rstrip() + addition + "\n", encoding="utf-8")


def create_file_manifest() -> None:
    rows = [
        ["path", "type", "related_id", "business_meaning", "agent_use"],
        ["files/excel/pipeline_obchodu_kveten_2026.xlsx", "xlsx", "deals", "pracovní obchodní pipeline od vedení", "extrahovat fáze obchodů, ceny, pravděpodobnost a další kroky"],
        ["files/excel/mesicni_report_leady_prodeje.xlsx", "xlsx", "monthly_metrics", "měsíční KPI report s grafy", "odpovědět na trend leadů a prodaných nemovitostí"],
        ["files/excel/data_quality_chybejici_podklady.xlsx", "xlsx", "properties", "seznam nemovitostí s nekompletními podklady", "najít blokery a doporučit další krok"],
        ["files/docx/tydenni_report_vedeni_2026-05-03.docx", "docx", "weekly_report", "report pro vedení", "vytěžit shrnutí, rizika a akční kroky"],
        ["files/docx/zapis_akvizicni_porada_2026-04-30.docx", "docx", "tasks", "zápis z interní porady", "převést rozhodnutí na úkoly"],
        ["files/pdf/list_vlastnictvi_p1001_mock.pdf", "pdf", "P-1001", "mock list vlastnictví", "ověřit vlastnictví a poznámky k nemovitosti"],
        ["files/pdf/penb_p1004_chybejici_vyzva.pdf", "pdf", "P-1004", "výzva k doložení PENB", "identifikovat compliance blocker"],
        ["files/contracts/rezervacni_smlouva_p1002_mock.pdf", "pdf", "P-1002", "pracovní rezervační smlouva", "najít otevřené body smlouvy"],
        ["google-calendar/google_calendar_events_mock.json", "json", "calendar", "napodobenina odpovědi Google Calendar API", "vybrat volný termín pro prohlídku"],
        ["google-calendar/pepa_google_calendar_mock.ics", "ics", "calendar", "export kalendáře", "fallback kalendářový import"],
    ]
    content = "\n".join(",".join(f'"{cell}"' if "," in cell else cell for cell in row) for row in rows) + "\n"
    (DEMO / "file_manifest.csv").write_text(content, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    create_excel_files()
    create_docx_files()
    create_pdf_files()
    create_google_calendar_files()
    create_file_manifest()
    update_readme()


if __name__ == "__main__":
    main()
